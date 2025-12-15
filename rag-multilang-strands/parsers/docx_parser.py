
from typing import List, Dict, Any
from io import BytesIO
from docx import Document
from parsers.common import table_to_html, flatten_table_text

def parse_docx_bytes(docx_bytes: bytes, fname: str) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    try:
        bio = BytesIO(docx_bytes)
        doc = Document(bio)

        paras = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                paras.append(p.text)
        if paras:
            out.append({"plain_text": "\n".join(paras), "source": fname})

        for ti, table in enumerate(doc.tables, start=1):
            headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
            rows = []
            for r in table.rows[1:]:
                rows.append([cell.text for cell in r.cells])

            if headers or rows:
                html = table_to_html(headers, rows)
                flat = flatten_table_text(headers, rows)
                out.append({
                    "table_html": html,
                    "plain_text": flat,
                    "source": f"{fname}#t{ti}"
                })
    except Exception as e:
        print(f"[parse_docx_bytes] Error parsing DOCX {fname}: {e}")
    return out

